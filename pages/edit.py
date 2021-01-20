# Inspired by Daniil "Zhinich" Ryaposov
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.oxml import OxmlElement, ns

def edit(document_name):
    document = Document(document_name)

    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.size = Pt(14)

    for section in document.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)

    for paragraph in document.paragraphs:
        paragraph.alignment = 3
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.5
        if 'graphicData' in paragraph._p.xml:
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.first_line_indent = Cm(0)
        else:
            paragraph.paragraph_format.first_line_indent = Cm(1.25)
            
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    paragraph.alignment = 1
                    paragraph.paragraph_format.line_spacing = 1
                    for run in paragraph.runs:
                        font = run.font
                        font.size = Pt(12)

    def create_element(name):
        return OxmlElement(name)

    def create_attribute(element, name, value):
        element.set(ns.qn(name), value)

    def add_page_number(run):
        fldChar1 = create_element('w:fldChar')
        create_attribute(fldChar1, 'w:fldCharType', 'begin')

        instrText = create_element('w:instrText')
        create_attribute(instrText, 'xml:space', 'preserve')
        instrText.text = "PAGE"

        fldChar2 = create_element('w:fldChar')
        create_attribute(fldChar2, 'w:fldCharType', 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

    page_number = document.sections[0].footer.paragraphs[0]
    page_number.clear()
    page_number.alignment = 1    
    add_page_number(page_number.add_run())   

    document.save(document_name)

